import logging
from pathlib import Path
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, RGBColor
from config import TEMPLATES_DIR, REPORTS_DIR

logger = logging.getLogger(__name__)


class DocxTemplateHandler:
    """Handle DOCX template management and filling."""

    def __init__(self):
        self.templates_dir = TEMPLATES_DIR
        self.reports_dir = REPORTS_DIR

    def get_available_templates(self) -> List[Dict[str, str]]:
        """Get list of available template files."""
        templates = []
        if self.templates_dir.exists():
            for template_file in self.templates_dir.glob("*.docx"):
                templates.append(
                    {
                        "name": template_file.stem,
                        "filename": template_file.name,
                        "path": str(template_file),
                    }
                )
        return templates

    def load_template(self, template_name: str) -> Document:
        """Load a template document."""
        template_path = self.templates_dir / f"{template_name}.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_name}")
        return Document(template_path)

    def replace_placeholder(
        self, doc: Document, placeholder: str, value: str
    ) -> None:
        """
        Replace placeholder text in document.
        Handles both direct text replacement and paragraph-level replacement.
        """
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                # Replace in runs while preserving formatting
                self._replace_text_in_runs(paragraph, placeholder, value)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            self._replace_text_in_runs(paragraph, placeholder, value)

    def _replace_text_in_runs(self, paragraph, placeholder: str, value: str) -> None:
        """Helper to replace text while preserving formatting."""
        if placeholder in paragraph.text:
            # Get the text and figure out how to replace it
            full_text = paragraph.text
            if placeholder in full_text:
                # Simple approach: clear and re-add with formatting from first run
                new_text = full_text.replace(placeholder, value)

                # Clear existing runs
                for run in paragraph.runs:
                    run.text = ""

                # Add new text in first run's style
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)

    def fill_finding(
        self,
        doc: Document,
        finding: Dict[str, Any],
        finding_number: int = 1,
    ) -> Document:
        """
        Fill in a finding in the document.
        Assumes template has placeholders like:
        [FINDING_TITLE], [DESCRIPTION_IMPACT], [REMEDIATION], [WSTG_REFERENCE]
        """
        placeholders = {
            "[FINDING_TITLE]": finding.get("title", "N/A"),
            "[DESCRIPTION_IMPACT]": finding.get("description_impact", "N/A"),
            "[REMEDIATION]": finding.get("remediation", "N/A"),
            "[WSTG_REFERENCE]": finding.get("wstg_reference", "N/A"),
        }

        for placeholder, value in placeholders.items():
            self.replace_placeholder(doc, placeholder, value)

        return doc

    def add_finding_section(
        self,
        doc: Document,
        finding: Dict[str, Any],
        finding_number: int,
    ) -> Document:
        """
        Add a finding as a new section (alternative to replacing placeholders).
        Useful if template has room for multiple findings.
        """
        # Add finding number
        heading = doc.add_heading(
            f"Finding {finding_number}: {finding.get('title', 'N/A')}", level=2
        )

        # Add description/impact
        doc.add_heading("Description & Impact", level=3)
        doc.add_paragraph(finding.get("description_impact", "N/A"))

        # Add remediation
        doc.add_heading("Remediation", level=3)
        doc.add_paragraph(finding.get("remediation", "N/A"))

        # Add WSTG reference
        if finding.get("wstg_reference"):
            doc.add_heading("Reference", level=3)
            p = doc.add_paragraph()
            run = p.add_run("WSTG Reference: ")
            ref_run = p.add_run(finding.get("wstg_reference", "N/A"))
            ref_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue for links

        return doc

    def save_report(self, doc: Document, filename: str) -> Path:
        """Save document to reports directory."""
        report_path = self.reports_dir / filename
        doc.save(report_path)
        logger.info(f"Report saved: {report_path}")
        return report_path

    def get_report_path(self, filename: str) -> Path:
        """Get path to a generated report."""
        return self.reports_dir / filename


def validate_template(template_path: Path) -> bool:
    """Validate that a DOCX file is a valid template."""
    try:
        doc = Document(template_path)
        # Check if it has placeholders
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text

        has_placeholders = "[FINDING_TITLE]" in full_text or "[DESCRIPTION_IMPACT]" in full_text
        return True  # Accept any valid DOCX for now
    except Exception as e:
        logger.error(f"Invalid template: {e}")
        return False
